"""
Document Processor - PyMuPDF4LLM based document processing
Optimal for PDF extraction with structure preservation
"""

import os
from typing import List, Dict, Any, Optional
from pathlib import Path
import logging

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Process documents using PyMuPDF4LLM for optimal extraction."""
    
    def __init__(self):
        self.supported_formats = {
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.txt': self._process_text,
            '.md': self._process_text,
        }
    
    def process_file(self, file_path: str) -> str:
        """Process file based on extension."""
        ext = Path(file_path).suffix.lower()
        
        if ext not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {ext}")
        
        processor = self.supported_formats[ext]
        return processor(file_path)
    
    def process_pdf(self, file_path: str) -> str:
        """Process PDF using PyMuPDF4LLM."""
        return self._process_pdf(file_path)
    
    def _process_pdf(self, file_path: str) -> str:
        """Extract text from PDF using PyMuPDF4LLM."""
        try:
            import pymupdf4llm
            
            # Extract text with page chunks
            md_text = pymupdf4llm.to_markdown(file_path)
            return md_text
        
        except ImportError:
            logger.warning("PyMuPDF4LLM not available, falling back to PyPDF2")
            return self._process_pdf_fallback(file_path)
        
        except Exception as e:
            logger.error(f"Error processing PDF: {e}")
            return self._process_pdf_fallback(file_path)
    
    def _process_pdf_fallback(self, file_path: str) -> str:
        """Fallback PDF processing using PyPDF2."""
        import PyPDF2
        
        text = ""
        try:
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
        except Exception as e:
            logger.error(f"Error in PDF fallback: {e}")
        
        return text
    
    def _process_docx(self, file_path: str) -> str:
        """Extract text from DOCX."""
        import docx
        
        text = ""
        try:
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    text += row_text + "\n"
        except Exception as e:
            logger.error(f"Error processing DOCX: {e}")
        
        return text
    
    def _process_text(self, file_path: str) -> str:
        """Process plain text files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            logger.error(f"Error reading text file: {e}")
            return ""
    
    def chunk_text(
        self,
        text: str,
        chunk_size: int = 512,
        overlap: int = 128
    ) -> List[Dict[str, Any]]:
        """Split text into overlapping chunks with metadata."""
        words = text.split()
        chunks = []
        
        for i in range(0, len(words), chunk_size - overlap):
            chunk_words = words[i:i + chunk_size]
            chunk_text = " ".join(chunk_words)
            
            chunks.append({
                "text": chunk_text,
                "word_count": len(chunk_words),
                "start_idx": i,
                "end_idx": min(i + chunk_size, len(words))
            })
        
        return chunks
    
    def extract_with_citations(
        self,
        file_path: str,
        chunk_size: int = 512,
        overlap: int = 128
    ) -> List[Dict[str, Any]]:
        """Extract text with page citations (for PDFs)."""
        ext = Path(file_path).suffix.lower()
        filename = Path(file_path).name
        
        if ext == '.pdf':
            return self._extract_pdf_with_pages(file_path, chunk_size, overlap)
        else:
            # Non-PDF: treat as single page
            text = self.process_file(file_path)
            chunks = self.chunk_text(text, chunk_size, overlap)
            
            for chunk in chunks:
                chunk["source"] = filename
                chunk["page"] = 1
            
            return chunks
    
    def _extract_pdf_with_pages(
        self,
        file_path: str,
        chunk_size: int,
        overlap: int
    ) -> List[Dict[str, Any]]:
        """Extract PDF text with page numbers."""
        import PyPDF2
        
        chunks = []
        filename = Path(file_path).name
        
        try:
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                
                for page_num, page in enumerate(reader.pages, 1):
                    page_text = page.extract_text()
                    if not page_text:
                        continue
                    
                    page_chunks = self.chunk_text(page_text, chunk_size, overlap)
                    
                    for chunk in page_chunks:
                        chunk["source"] = filename
                        chunk["page"] = page_num
                        chunks.append(chunk)
        
        except Exception as e:
            logger.error(f"Error extracting PDF with pages: {e}")
        
        return chunks